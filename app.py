import streamlit as st
import pandas as pd
from docx import Document
import zipfile
import os
import shutil

st.title("📄 Appointment Letter Generator")

# File uploaders for Excel and Word template
excel_file = st.file_uploader("📂 Upload Excel File", type=["xls", "xlsx"])
docx_file = st.file_uploader("📂 Upload Word Template", type=["docx"])

output_folder = "generated_letters"

# Function to replace text placeholders in paragraphs
def replace_text_in_paragraphs(doc, replacements):
    for para in doc.paragraphs:
        for key, value in replacements.items():
            if f"{{{key}}}" in para.text:
                para.text = para.text.replace(f"{{{key}}}", str(value))

# Function to replace text inside tables
def replace_text_in_table(doc, replacements):
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if f"{{{key}}}" in cell.text:
                        cell.text = cell.text.replace(f"{{{key}}}", str(value))

# Function to generate appointment letters
def generate_letter(template_docx, data, output_file):
    doc = Document(template_docx)
    replace_text_in_paragraphs(doc, data)
    replace_text_in_table(doc, data)
    doc.save(output_file)

if excel_file and docx_file:
    df = pd.read_excel(excel_file)

    # Check if required columns are present
    required_columns = {
        "Name", "EmpCode", "Address", "Designation", "Department", "Date",
        "Basic Monthly", "Basic Annual", "HRA Monthly", "HRA Annual", 
        "Statutory Bonus Monthly", "Statutory Bonus Annual", 
        "Special Allowance Monthly", "Special Allowance Annual", 
        "Gross Salary Monthly", "Gross Salary Annual",
        "Employer PF Monthly", "Employer PF Annual", "Employer LWF Monthly", "Employer LWF Annual",
        "Employer ESIC Monthly", "Employer ESIC Annual", "Cost to Company Monthly", "Cost to Company Annual",
        "Employee PF Monthly", "Employee PF Annual", "Employee LWF Monthly", "Employee LWF Annual",
        "Employee ESIC Monthly", "Employee ESIC Annual", "Total Deduction Monthly", "Total Deduction Annual",
        "Net Take Home Monthly", "Net Take Home Annual",
        "Statutory Leave Monthly", "Statutory Leave Annual",
        "Ex gratia Pay Monthly", "Ex gratia Pay Annual"
    }

    if not required_columns.issubset(df.columns):
        st.error("❌ The Excel file must contain all required columns.")
    else:
        st.success("✅ Excel and Word template successfully uploaded!")

        if st.button("Generate and Download ZIP"):
            # Create output folder
            if not os.path.exists(output_folder):
                os.makedirs(output_folder)

            file_list = []
            for _, row in df.iterrows():
                data_dict = {
                    "input_name": row["Name"],
                    "input_empcode": row["EmpCode"],
                    "input_address": row["Address"],
                    "input_designation": row["Designation"],
                    "input_department": row["Department"],
                    "input_date": row["Date"].strftime("%d-%m-%Y") if pd.notna(row["Date"]) else "",

                    # Monthly & Annual Salary Breakdown
                    "input_basic_monthly": row["Basic Monthly"],
                    "input_basic_annual": row["Basic Annual"],
                    "input_hra_monthly": row["HRA Monthly"],
                    "input_hra_annual": row["HRA Annual"],
                    "input_bonus_monthly": row["Statutory Bonus Monthly"],
                    "input_bonus_annual": row["Statutory Bonus Annual"],
                    "input_special_allowance_monthly": row["Special Allowance Monthly"],
                    "input_special_allowance_annual": row["Special Allowance Annual"],
                    "input_gross_salary_monthly": row["Gross Salary Monthly"],
                    "input_gross_salary_annual": row["Gross Salary Annual"],

                    # Employer Contributions
                    "input_employer_pf_monthly": row["Employer PF Monthly"],
                    "input_employer_pf_annual": row["Employer PF Annual"],
                    "input_employer_lwf_monthly": row["Employer LWF Monthly"],
                    "input_employer_lwf_annual": row["Employer LWF Annual"],
                    "input_employer_esic_monthly": row["Employer ESIC Monthly"],
                    "input_employer_esic_annual": row["Employer ESIC Annual"],
                    "input_ctc_monthly": row["Cost to Company Monthly"],
                    "input_ctc_annual": row["Cost to Company Annual"],

                    # Deductions
                    "input_employee_pf_monthly": row["Employee PF Monthly"],
                    "input_employee_pf_annual": row["Employee PF Annual"],
                    "input_employee_lwf_monthly": row["Employee LWF Monthly"],
                    "input_employee_lwf_annual": row["Employee LWF Annual"],
                    "input_employee_esic_monthly": row["Employee ESIC Monthly"],
                    "input_employee_esic_annual": row["Employee ESIC Annual"],
                    "input_total_deduction_monthly": row["Total Deduction Monthly"],
                    "input_total_deduction_annual": row["Total Deduction Annual"],
                    "input_net_take_home_monthly": row["Net Take Home Monthly"],
                    "input_net_take_home_annual": row["Net Take Home Annual"],

                    # New Pay Heads
                    "input_statutory_leave_monthly": row["Statutory Leave Monthly"],
                    "input_statutory_leave_annual": row["Statutory Leave Annual"],
                    "input_exgratia_monthly": row["Ex gratia Pay Monthly"],
                    "input_exgratia_annual": row["Ex gratia Pay Annual"],
                }

                file_name = f"Appointment_Letter_{row['Name']}_{row['EmpCode']}.docx"
                file_path = os.path.join(output_folder, file_name)

                generate_letter(docx_file, data_dict, file_path)
                file_list.append(file_path)

            # Create a ZIP file
            zip_filename = "Appointment_Letters.zip"
            shutil.make_archive("Appointment_Letters", "zip", output_folder)

            # Provide a download button for the ZIP file
            with open(zip_filename, "rb") as zipf:
                st.download_button(
                    label="Download All Letters (ZIP)",
                    data=zipf,
                    file_name="Appointment_Letters.zip",
                    mime="application/zip"
                )

            st.success("✅ All letters generated successfully!")
